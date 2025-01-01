import os
import shutil
import sys
import uuid

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from xlutils.copy import copy

import pandas as pd
import requests
import xlrd
import xlwt
from PIL import Image
from io import BytesIO
from docxtpl import DocxTemplate

from add_float_picture import add_float_picture
from mail_merge import mail_merge


def column_letter_to_index(letter):
    """将 Excel 列字母转换为数字索引（从0开始）"""
    letter = letter.upper()
    index = 0
    for char in letter:
        index = index * 26 + (ord(char) - ord('A') + 1)
    return index - 1


def generate_randa(uid, user_name, class_name, dimension, data):
    # 定义请求的 URL
    # url = 'https://fivespace.vercel.app/api/chart'  # 替换为你的服务器地址
    url = 'http://123.60.83.109:9988/api/chart'  # 替换为你的服务器地址

    # 定义请求的 JSON 数据
    # data = {
    #     "width": 1024,  # 指定图表宽度
    #     "height": 768,  # 指定图表高度
    #     "option": {
    #         "title": {
    #             "text": "五维评价"
    #         },
    #         "legend": {
    #             "data": ["张三", "平均分"]
    #         },
    #         "radar": {
    #             "indicator": [
    #                 {"name": "爱国情怀", "max": 6500},
    #                 {"name": "正直诚信", "max": 16000},
    #                 {"name": "遵规守纪", "max": 30000},
    #                 {"name": "仁爱友善", "max": 38000},
    #                 {"name": "包容理解", "max": 52000}
    #             ]
    #         },
    #         "series": [
    #             {
    #                 "name": "Budget vs spending",
    #                 "type": "radar",
    #                 "data": [
    #                     {
    #                         "value": [4200, 3000, 20000, 35000, 50000],
    #                         "name": "张三"
    #                     },
    #                     {
    #                         "value": [5000, 14000, 28000, 26000, 42000],
    #                         "name": "平均分"
    #                     }
    #                 ]
    #             }
    #         ]
    #     }
    # }

    # 发送 POST 请求
    response = requests.post(url, json=data)

    # 检查请求是否成功
    if response.status_code == 200:
        # 将响应的二进制数据转换为图像
        image = Image.open(BytesIO(response.content))
        # 显示图像
        # image.show()
        # 或者保存图像到本地文件
        png_path = f'./pic/{uid}-{dimension}.png'
        image.save(png_path)
        return png_path

    else:
        print(f"请求失败，状态码: {response.status_code}")
        return ''


# def test_xls(file_path):
#     # 打开 .xls 文件进行读取
#     workbook = xlrd.open_workbook(file_path)
#     sheet = workbook.sheet_by_name('1')  # 假设要读取的工作表名为 'Sheet1'
#
#     # 获取最后一列的索引
#     last_col_index = sheet.ncols - 1
#
#
#     # 读取指定列的数据并写入新文件
#     for row_idx in range(sheet.nrows):
#         value = sheet.cell(row_idx, column_index).value
#         new_sheet.write(row_idx, 0, value)  # 将数据写入新文件的第一列
#
#     # 保存到新的 .xls 文件
#     new_workbook.save(output_path)

def generate_randa1(file_path, dest_sheet, column_data):
    # 读取整个工作表
    df = pd.read_excel(file_path, sheet_name='1')

    # 检查并创建 'uid' 列，如果它不存在
    if 'uid' not in df.columns:
        df['uid'] = str(uuid.uuid4())  # 或者 df['uid'] = '' / df['uid'] = 0, 根据需要初始化

    # 选择第2个工作表
    df2 = pd.read_excel(file_path, sheet_name='基础数据录入')

    # 提取特定单元格的值
    school_year = df2.iloc[0, 1]  # 获取第1行第2列的值
    semester = df2.iloc[1, 1]  # 获取第2行第2列的值
    class_info = df2.iloc[2, 1]  # 获取第3行第2列的值
    teacher = df2.iloc[3, 1]  # 获取第4行第2列的值

    # 打印提取的值
    print("学年:", school_year)
    print("学期:", semester)
    print("班级:", class_info)
    print("班主任:", teacher)

    # df_cleaned = df.dropna(subset=['姓名'])  # 去掉 '姓名' 列中为空的行

    # 定义每个维度的列名
    columns_map = {
        "正": ["爱国情怀", "正直诚信", "遵规守纪", "仁爱友善", "包容理解"],
        "善": ["学习态度", "学业进步", "学科成绩", "阅读习惯", "科学探索"],
        "健": ["体质体能", "运动技能", "体锻习惯", "健康生活", "课堂表现"],
        "雅": ["文明礼仪", "文化修养", "审美情趣", "音乐成绩", "美术成绩"],
        "勤": ["劳动意识", "劳动技能", "劳动实践", "课程表现", "劳动创新"]
    }

    # 计算每个维度的平均分
    average_scores_map = {dimension: df[columns].mean().tolist() for dimension, columns in columns_map.items()}

    # # 为每个维度创建新的列来存储雷达图路径
    # for dimension in columns_map.keys():
    #     df_cleaned.loc[:, f'{dimension}'] = None
    # 为每个维度创建新的列来存储雷达图路径
    # 为每个维度创建新的列来存储雷达图路径
    for dimension in columns_map.keys():
        df[f'{dimension}'] = pd.Series([""] * len(df), dtype="str")

    # 遍历每个学生，生成对应的雷达图数据
    for index, row in df.iterrows():
        student_name = row['姓名']
        uid = row['uid']
        if pd.isna(uid):
            uid = str(uuid.uuid4())

        student_scores_map = {dimension: row[columns].tolist() for dimension, columns in columns_map.items()}

        for dimension, student_scores in student_scores_map.items():
            average_scores = average_scores_map[dimension]
            print(f"{student_name} - {dimension}维度分数: {student_scores}")
            print(f"{dimension}维度平均分: {average_scores}")

            # 动态生成 indicator 列表
            indicators = [{"name": name, "max": 10} for name in columns_map[dimension]]

            data = {
                "width": 500,  # 指定图表宽度
                "height": 500,  # 指定图表高度
                "option": {
                    "title": {
                        "text": f""
                    },
                    "legend": {
                        # "data": [student_name, "平均分"]
                        "data": ['平均', '我的']
                    },
                    "radar": {
                        "indicator": indicators
                        # "indicator": [
                        #     {"name": "爱国情怀", "max": 10},
                        #     {"name": "正直诚信", "max": 10},
                        #     {"name": "遵规守纪", "max": 10},
                        #     {"name": "仁爱友善", "max": 10},
                        #     {"name": "包容理解", "max": 10}
                        # ]
                    },
                    "series": [
                        {
                            "name": "五维评价",
                            "type": "radar",
                            "data": [
                                {
                                    "value": student_scores,
                                    "name": "我的"  # student_name
                                },
                                {
                                    "value": average_scores,
                                    "name": "平均"
                                }
                            ]
                        }
                    ]
                }
            }
            randa_path = generate_randa(uid, student_name, class_info, dimension, data)
            # 修改特定单元格的值
            column_index = column_letter_to_index(column_data[dimension])
            column_uid_index = df.columns.get_loc("uid")

            # column_uid_index = column_letter_to_index("uid")
            dest_sheet.write(index + 1, column_index, randa_path)
            dest_sheet.write(index + 1, column_uid_index, str(uid))

            # 将路径存储到数据框中
            # df.at[index, f'{dimension}'] = randa_path
            # break
    # df.to_excel('./output/五维评价雷达图.xlsx', index=False)
    # 使用 ExcelWriter 保存修改
    # with pd.ExcelWriter(file_path, engine='openpyxl', mode='a', if_sheet_exists='replace') as writer:
    #     df.to_excel(writer, sheet_name='Sheet1', index=False)
    # 保存修改后的数据到新的 .xls 文件
    # df.to_excel(file_path+".1.xls", sheet_name='Sheet1', index=False, engine='xlwt')


def parse_excel(file_path):
    # 读取指定工作表和列
    df = pd.read_excel(file_path, sheet_name='1', usecols='B,W:AA')
    # 去掉 B 列中为空或者 NaN 的行
    df_cleaned = df.dropna(subset=['姓名'])  # 使用 'B' 列名来检查空值

    # print(df_cleaned)
    # 计算每个维度的平均分
    average_scores = df_cleaned.iloc[:, 1:].mean().tolist()

    # 遍历每个学生，生成对应的雷达图数据
    for index, row in df_cleaned.iterrows():
        student_name = row['姓名']
        student_scores = row.iloc[1:].tolist()
        print(student_name, student_scores)

        # 准备雷达图数据结构
        data = {
            "width": 300,  # 指定图表宽度
            "height": 300,  # 指定图表高度
            "option": {
                "title": {
                    "text": f"{student_name}的五维评价"
                },
                "legend": {
                    "data": [student_name, "平均分"]
                },
                "radar": {
                    "indicator": [
                        {"name": "爱国情怀", "max": 10},
                        {"name": "正直诚信", "max": 10},
                        {"name": "遵规守纪", "max": 10},
                        {"name": "仁爱友善", "max": 10},
                        {"name": "包容理解", "max": 10}
                    ]
                },
                "series": [
                    {
                        "name": "个人 vs 平均",
                        "type": "radar",
                        "data": [
                            {
                                "value": student_scores,
                                "name": student_name
                            },
                            {
                                "value": average_scores,
                                "name": "平均分"
                            }
                        ]
                    }
                ]
            }
        }
        # print(data)
        generate_randa(student_name, data)
    # sys.exit(0)


from docx import Document
import re


def _replace_paragraphs(doc, replacements, df, sheet_row):
    for paragraph in doc.paragraphs:
        for replace_item in replacements:
            key = f"«{replace_item}»"
            value = sheet_row[replace_item]
            if key in paragraph.text:
                # if replace_item == '姓名':
                #     print("ss")
                value = sheet_row[replace_item]
                # 如果值是 NaN，则替换为 ''
                if pd.isna(value):
                    value = ''
                else:
                    # if pd.api.types.is_numeric_dtype(df[replace_item]):
                    #     value = str(int(value))  # 用 0 替换数值型 NaN
                    # else:
                    value = str(value)  # 用空字符串替换非数值型 NaN
                paragraph.text = paragraph.text.replace(key, value)


def _replace_table(doc, replacements, df, sheet_row):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # print("=====",cell.text)
                # if '学生签字' in cell.text:
                #     print(cell.text)
                for replace_item in replacements:
                    key = f"«{replace_item}»"
                    value = sheet_row[replace_item]
                    # 如果值是 NaN，则替换为 ''
                    if pd.isna(value):
                        value = ''
                    else:
                        value = str(value)
                        # if pd.api.types.is_numeric_dtype(df[replace_item]):
                        #     value = str(int(value))  # 用 0 替换数值型 NaN
                        # else:
                        #     value = str(value)  # 用空字符串替换非数值型 NaN
                    if key in cell.text:
                        # if '学生签字' in cell.text:
                        #     print(cell.text)
                        cell.text = cell.text.replace(key, value)
                # for replace_item in replacements:
                #     key = f"«{replace_item}»"
                #     value = sheet_row[replace_item]


def replace_text_in_docx(excel_path, doc_path):
    # tpl = DocxTemplate(doc_path)
    # print(tpl.get_document_xml())

    df = pd.read_excel(excel_path, sheet_name='1')
    replacements = ["50米×8往返跑", "50米×8往返跑等级", "50米×8往返跑评分", "50米跑", "50米跑等级", "50米跑评分",
                    "A体重", "A身高", "B视力右", "B视力左", "X收缩压", "X脉搏", "X舒张压", "一分钟仰卧起坐",
                    "一分钟仰卧起坐等级", "一分钟仰卧起坐评分", "一分钟跳绳", "一分钟跳绳等级", "一分钟跳绳评分",
                    "书法", "事假", "仁爱友善", "体育", "体质体能", "体锻习惯", "作文", "信息", "健康生活", "劳动",
                    "劳动创新", "劳动实践", "劳动意识", "劳动技能", "包容理解", "听力右", "听力左", "坐位体前屈",
                    "坐位体前屈等级", "坐位体前屈评分", "姓名", "学业进步", "学习态度", "学科成绩", "审美情趣", "总分",
                    "总分等级", "总评", "操行等级", "数学", "文化修养", "文明礼仪", "旷课", "正直诚信", "爱国情怀",
                    "班级职务", "病假", "科学", "科学探索", "综合实践", "美术", "美术成绩", "肺活量", "英语", "评语",
                    "课堂表现", "课程表现", "运动技能", "迟到", "道德与法治", "遵规守纪", "阅读", "阅读习惯", "附加分",
                    "音乐", "音乐成绩", "龋齿"]
    # student_name  = df["姓名"]

    # 遍历 DataFrame 的每一行
    for index, sheet_row in df.iterrows():
        # 为每一行创建一个新的文档对象
        new_doc = Document(doc_path)  # 从原始模板创建
        student_name = sheet_row['姓名']
        for field in replacements:
            print(f"{field}: {sheet_row[field]}")

        _replace_paragraphs(new_doc, replacements, df, sheet_row)

        _replace_table(new_doc, replacements, df, sheet_row)

        # 保存每一条记录到新的 Word 文件
        # new_output_path = output_path.replace('.docx', f'_{index}.docx')
        new_output_path = f"./output/{student_name}.docx"
        new_doc.save(new_output_path)
        print(f"Saved: {new_output_path}")
        sys.exit(0)


def generate_five_pic(src, destination_path):
    # shutil.copy(src, destination_path)
    # workbook = xlrd.open_workbook(destination_path)
    # dest_sheet = workbook.sheet_by_name('1')  # 假设要读取的工作表名为 'Sheet1'
    data = xlrd.open_workbook(src, formatting_info=True)

    excel = copy(wb=data)  # 完成xlrd对象向xlwt对象转换
    dest_sheet = excel.get_sheet("1")  # 获得要操作的页

    column_data = {
        "正": "CL",
        "善": "CM",
        "健": "CN",
        "雅": "CO",
        "勤": "CP",

    }
    # column_letter_to_index("CL")

    generate_randa1(src, dest_sheet, column_data)
    excel.save(destination_path)


from docx import Document
from docx.shared import Inches, Cm


def replace_placeholder_with_image(doc, placeholder, image_path, width=None, height=None):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            # Create a new run with the image
            new_run = paragraph.add_run()
            new_run.add_picture(image_path, width=width, height=height)

            # Remove the placeholder
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, '')

            break


# 替换图片
def replace_image(run, new_image_path):
    # 生成新的图片ID
    rels = run.part.rels
    old_rId = rels.rel_id
    new_rId = rels.get_or_add_relationship(
        reltype=rels[old_rId].reltype,
        target=new_image_path
    ).rel_id

    # 替换图片
    drawing = run._r.get_or_add_drawing()
    inline = drawing.inline
    graphic = inline.graphic
    graphicData = graphic.graphicData
    pic = graphicData.pic
    blipFill = pic.blipFill
    blip = blipFill.blip
    blip.embed = new_rId


def enum_img():
    image_path = 'D:\\workspace\\xhp\\schoolmgr\\server\\scripts\\fivespace\\pic\\德-一年（2）班-胡欣橙.png'
    # Load an existing document
    new_doc = Document("./data/Doc1.docx")
    # 替换图片
    for para in new_doc.paragraphs:
        for run in para.runs:
            replace_image(run, image_path)
            # print(run._r.xml)
            # if run._r.xml.endswith('</w:drawing>'):
            #     print('11')
            #     # replace_image(run, image_path)
    new_doc.save('output.docx')

    # Replace the placeholder with the image
    # replace_placeholder_with_image(doc, '<<image_placeholder>>', 'image.png', width=Inches(1.5), height=Inches(1.5))
    #
    # # Save the document
    # doc.save('output.docx')


def test_template_pic():
    from docxtpl import DocxTemplate

    # 初始化DocxTemplate类的对象，用于后续的模板渲染
    # 参数为模板文件的路径
    # tpl = DocxTemplate('./data/素质评价单（两栏装订）.docx')
    tpl = DocxTemplate('./data/素质评价单（两栏装订）.docx')

    # 创建一个空字典，用于存储模板渲染时所需的上下文数据
    context = {}

    # 使用上下文数据渲染模板
    # 这里的render方法会根据context中的数据替换模板中对应的部分
    tpl.render(context)

    # 使用xpath方法查找docx文件中所有图片的属性
    # 图片的属性存储在pic:cNvPr节点中
    pic_attrib_list = [pic.attrib for pic in tpl.docx.element.xpath(".//pic:cNvPr")]

    # 遍历所有图片的属性并打印
    # 这里的目的是展示或处理模板中的图片属性
    for attr in pic_attrib_list:
        print(attr)


def replace_pic():
    image_path = 'D:\\workspace\\xhp\\schoolmgr\\server\\scripts\\fivespace\\pic\\德-一年（2）班-胡欣橙.png'
    tpl = DocxTemplate('./data/Doc1.docx')
    context = {}
    tpl.render(context)
    tpl.replace_pic('1234', image_path)
    tpl.save('./test.docx')

def replace_pic_with_excel(excel_path):
    df = pd.read_excel(excel_path, sheet_name='1').fillna('')
    for index, row in df.iterrows():
        # student_name = row['姓名']
        uid = row['uid']

        dict = [
            {'正01': row['正']},
            {'善01': row['善']},
            {'健01': row['健']},
            {'雅01': row['雅']},
            {'勤01': row['勤']},
        ]
        # 打开word 替换
        src = f'./output/{uid}.docx'
        dest = f'./output/final_{uid}.docx'
        tpl = DocxTemplate(src)
        context = {}
        tpl.render(context)
        for item in dict:
            for key, value in item.items():
                tpl.replace_pic(key, value)

        tpl.save(dest)


if __name__ == '__main__':
    src = "./data/全科成绩及五育评价数据录入表1.xls"
    template = './data/素质评价单（两栏装订）.docx'
    destination_path = src + "randa.xls"
    # 生成五维图
    generate_five_pic(src, destination_path)

    # 批量合并邮件
    mail_merge(destination_path, template)

    # 替换图片
    replace_pic_with_excel(destination_path)
